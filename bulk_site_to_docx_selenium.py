"""
BALANCED PRO EXTRACTOR v3
Fixes: hidden accordion/FAQ/tab/slider content via JS force-reveal
Elementor + Non-Elementor Safe
"""

import json
import time
import requests
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==============================
# USER SETTINGS
# ==============================

BASE_URL = "https://localrooferdev.wpenginepowered.com/"
WAIT_TIME = 3
MAX_RETRIES = 3

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7)"
}

LOW_PRIORITY_PATTERNS = [
    "/category/",
    "/tag/",
    "/author/",
    "/hello-world",
]

CTA_PHRASES = [
    "call now",
    "get in touch",
    "schedule",
    "request",
    "learn more",
    "submit",
    "free inspection",
    "contact us",
]

# ==============================
# SELENIUM SETUP
# ==============================

chrome_options = Options()
chrome_options.add_argument("--headless=new")
chrome_options.add_argument("--disable-gpu")
chrome_options.add_argument("--no-sandbox")
chrome_options.add_argument("--window-size=1920,1080")       # smaller height reduces memory per tab
chrome_options.add_argument("--disable-dev-shm-usage")       # use /tmp instead of /dev/shm (fixes WSL2 tab crashes)
chrome_options.add_argument("--disable-extensions")
chrome_options.add_argument("--disable-software-rasterizer")
chrome_options.add_argument("--js-flags=--max-old-space-size=512")  # cap JS heap at 512 MB
chrome_options.add_argument("--disable-background-networking")
chrome_options.add_argument("--disable-default-apps")
chrome_options.add_argument("--disable-sync")
chrome_options.add_argument("--metrics-recording-only")
chrome_options.add_argument("--mute-audio")
chrome_options.add_argument("--no-first-run")
chrome_options.add_argument("--safebrowsing-disable-auto-update")

def make_driver():
    return webdriver.Chrome(options=chrome_options)

print("Starting browser...")
driver = make_driver()

# ==============================
# SITEMAP
# ==============================

def find_sitemap_urls(base_url):
    candidates = [
        "/sitemap.xml",
        "/sitemap_index.xml",
        "/wp-sitemap.xml",
    ]

    found_urls = []

    for path in candidates:
        url = urljoin(base_url, path)
        try:
            r = requests.get(url, timeout=20, headers=HEADERS)
            if r.status_code == 200:
                txt = r.text.lower()
                if "<urlset" in txt or "<sitemapindex" in txt:
                    found_urls.append(url)
        except Exception:
            pass

    return list(set(found_urls))


def parse_sitemap(url, collected):
    try:
        r = requests.get(url, timeout=30, headers=HEADERS)
        soup = BeautifulSoup(r.text, "xml")

        sitemap_tags = soup.find_all("sitemap")
        if sitemap_tags:
            for sm in sitemap_tags:
                parse_sitemap(sm.loc.text.strip(), collected)
            return

        for url_tag in soup.find_all("url"):
            collected.add(url_tag.loc.text.strip())

    except Exception as e:
        print(f"Sitemap parse error: {e}")

# ==============================
# PAGE INTERACTION
# ==============================

def smart_scroll():
    try:
        for _ in range(14):
            driver.execute_script("window.scrollBy(0, window.innerHeight);")
            time.sleep(0.6)
    except Exception:
        pass


def force_reveal_hidden_content():
    """
    Force all typically-hidden content panels to be visible via JavaScript.
    Run this AFTER scrolling and BEFORE parsing page_source.
    """
    try:
        driver.execute_script("""
            // Pass 1: Accordion / FAQ / Collapse panels
            var accordionSelectors = [
                '.elementor-tab-content',
                '.elementor-accordion-content',
                '.elementor-toggle-content',
                '[class*="accordion-content"]',
                '[class*="faq-answer"]',
                '[class*="faq-content"]',
                '[class*="collapse"]',
                '.e-n-accordion-item__body'
            ];
            accordionSelectors.forEach(function(sel) {
                document.querySelectorAll(sel).forEach(function(el) {
                    el.style.display = 'block';
                    el.style.height = 'auto';
                    el.style.maxHeight = 'none';
                    el.style.overflow = 'visible';
                    el.style.visibility = 'visible';
                    el.style.opacity = '1';
                    el.removeAttribute('aria-hidden');
                    el.removeAttribute('hidden');
                });
            });

            // Pass 2: Tab panels
            var tabSelectors = [
                '[role="tabpanel"]',
                '.tab-pane',
                '.ui-tabs-panel',
                '.elementor-tab-content'
            ];
            tabSelectors.forEach(function(sel) {
                document.querySelectorAll(sel).forEach(function(el) {
                    el.style.display = 'block';
                    el.classList.remove('hidden', 'inactive');
                });
            });

            // Pass 3: Force ALL swiper/carousel slides fully visible (including clones)
            document.querySelectorAll('.swiper-slide, .slick-slide, .owl-item, .splide__slide, .elementor-carousel-item').forEach(function(el) {
                el.style.display = 'block';
                el.style.visibility = 'visible';
                el.style.opacity = '1';
                el.removeAttribute('aria-hidden');
                el.removeAttribute('hidden');
                el.removeAttribute('inert');
            });

            // Pass 4: Generic inline-hidden content-bearing elements
            var excludeTags = ['NAV', 'HEADER', 'FOOTER'];
            var excludeClasses = ['modal', 'overlay', 'popup', 'cookie', 'banner'];
            var hiddenEls = document.querySelectorAll(
                '[style*="display: none"], [style*="display:none"], [hidden], [aria-hidden="true"], [inert]'
            );
            hiddenEls.forEach(function(el) {
                if (excludeTags.indexOf(el.tagName) !== -1) return;
                var cls = (el.className && typeof el.className === 'string') ? el.className : '';
                for (var i = 0; i < excludeClasses.length; i++) {
                    if (cls.indexOf(excludeClasses[i]) !== -1) return;
                }
                el.style.display = 'block';
                el.style.height = 'auto';
                el.style.maxHeight = 'none';
                el.style.overflow = 'visible';
                el.style.visibility = 'visible';
                el.style.opacity = '1';
                el.removeAttribute('aria-hidden');
                el.removeAttribute('hidden');
                el.removeAttribute('inert');
            });
        """)
        time.sleep(0.3)
    except Exception:
        pass


def expand_dynamic_content():
    """
    Click accordion-specific triggers only.
    Generic 'button' selector removed to prevent form submits / modal triggers.
    Runs 2 rounds; JS force-reveal handles the rest.
    """
    try:
        selectors = [
            ".elementor-tab-title",
            ".elementor-accordion-title",
            ".elementor-toggle-title",
            "[role='tab']",
            "summary",
        ]

        for _ in range(2):
            for sel in selectors:
                elements = driver.find_elements(By.CSS_SELECTOR, sel)
                for el in elements:
                    try:
                        driver.execute_script(
                            "arguments[0].scrollIntoView({block:'center'});", el
                        )
                        driver.execute_script("arguments[0].click();", el)
                        time.sleep(0.25)
                    except Exception:
                        pass
    except Exception:
        pass


def activate_sliders():
    """
    Cycle through all slide next-arrows (20 clicks each),
    then call force_reveal_hidden_content() to ensure all slide content is visible.
    """
    try:
        arrows = driver.find_elements(
            By.CSS_SELECTOR,
            ".swiper-button-next, .slick-next, .elementor-swiper-button-next, "
            ".owl-next, .splide__arrow--next, [aria-label='Next']",
        )
        for arrow in arrows:
            for _ in range(20):
                try:
                    driver.execute_script("arguments[0].click();", arrow)
                    time.sleep(0.4)
                except Exception:
                    pass

        # After cycling slides, force all content visible
        force_reveal_hidden_content()

    except Exception:
        pass

# ==============================
# ⭐ FAQ SPECIAL HANDLER
# ==============================

def extract_elementor_faq(soup):
    faq_blocks = []
    seen_questions = set()

    # Path 1: Classic Elementor accordion
    for item in soup.select(".elementor-accordion-item"):
        q = item.select_one(".elementor-tab-title")
        a = item.select_one(".elementor-tab-content")
        if q and a:
            question = q.get_text(" ", strip=True)
            answer = a.get_text(" ", strip=True)
            if question and answer and question.lower() not in seen_questions:
                seen_questions.add(question.lower())
                faq_blocks.append(("h3", question))
                faq_blocks.append(("p", answer))

    # Path 2: Elementor 3.x nested accordion
    for item in soup.select(".e-n-accordion-item"):
        q = item.select_one(".e-n-accordion-item-title-text") or item.select_one("summary")
        a = item.select_one(".e-n-accordion-item__body")
        if q and a:
            question = q.get_text(" ", strip=True)
            answer = a.get_text(" ", strip=True)
            if question and answer and question.lower() not in seen_questions:
                seen_questions.add(question.lower())
                faq_blocks.append(("h3", question))
                faq_blocks.append(("p", answer))

    # Path 3: Native HTML5 <details>/<summary> pairs (non-destructive)
    for detail in soup.find_all("details"):
        summary_tag = detail.find("summary")
        if not summary_tag:
            continue
        question = summary_tag.get_text(" ", strip=True)
        answer_parts = []
        for child in detail.children:
            if child == summary_tag:
                continue
            if hasattr(child, "get_text"):
                part = child.get_text(" ", strip=True)
                if part:
                    answer_parts.append(part)
        answer = " ".join(answer_parts)
        if question and answer and question.lower() not in seen_questions:
            seen_questions.add(question.lower())
            faq_blocks.append(("h3", question))
            faq_blocks.append(("p", answer))

    # Path 4: Schema.org FAQ JSON-LD
    for script in soup.find_all("script", {"type": "application/ld+json"}):
        try:
            data = json.loads(script.string or "")
            if isinstance(data, dict):
                data = [data]
            for obj in data:
                if obj.get("@type") == "FAQPage":
                    for entity in obj.get("mainEntity", []):
                        q = entity.get("name", "")
                        a_obj = entity.get("acceptedAnswer", {})
                        a = a_obj.get("text", "") if isinstance(a_obj, dict) else ""
                        if q and a and q.lower() not in seen_questions:
                            seen_questions.add(q.lower())
                            a_clean = BeautifulSoup(a, "lxml").get_text(" ", strip=True)
                            faq_blocks.append(("h3", q))
                            faq_blocks.append(("p", a_clean))
        except Exception:
            pass

    return faq_blocks

# ==============================
# CLEAN TEXT
# ==============================

def clean_text_blocks(soup):

    for tag in soup(["script","style","noscript","header","footer","nav","svg"]):
        tag.decompose()

    # Remove only form input/control elements, not the form wrapper itself
    # (the form container may hold headings and text like "GET IN TOUCH")
    for tag in soup(["input","textarea","select","option","label"]):
        tag.decompose()

    text_items = []
    seen = set()        # lowercased text strings already emitted
    seen_nodes = set()  # id() of accordion item nodes already emitted



    def _has_class(el, cls):
        """Check if a BeautifulSoup element has a specific class (among possibly many)."""
        return cls in (el.get("class") or [])

    def _find_accordion_parent(el):
        """Walk up the tree looking for an accordion item node."""
        for parent in el.parents:
            if _has_class(parent, "elementor-accordion-item") or _has_class(parent, "e-n-accordion-item"):
                return parent
        return None

    # Single unified walk in DOM order — preserves exact page sequence.
    # We walk ALL tags so we can catch text-editor divs with bare text nodes.
    root = soup.body if soup.body else soup
    for el in root.find_all(True, recursive=True):

        # ── Accordion items: emit full Q/A pair inline, skip all their children ──
        if _has_class(el, "elementor-accordion-item") or _has_class(el, "e-n-accordion-item"):
            acc_id = id(el)
            if acc_id not in seen_nodes:
                seen_nodes.add(acc_id)

                # Classic Elementor accordion
                q_tag = el.select_one(".elementor-tab-title")
                a_tag = el.select_one(".elementor-tab-content")

                # Elementor 3.x nested accordion (<details>/<summary> structure)
                # Question: .e-n-accordion-item-title-text inside <summary>
                # Answer: <div role="region"> after <summary> (contains text-editor widget)
                if not q_tag:
                    q_tag = el.select_one(".e-n-accordion-item-title-text")
                if not a_tag:
                    a_tag = el.find("div", attrs={"role": "region"})

                if q_tag and a_tag:
                    q_text = q_tag.get_text(" ", strip=True)
                    a_text = a_tag.get_text(" ", strip=True)
                    if q_text and q_text.lower() not in seen:
                        seen.add(q_text.lower())
                        text_items.append(("h3", q_text))
                    if a_text and a_text.lower() not in seen:
                        seen.add(a_text.lower())
                        text_items.append(("p", a_text))
            continue

        # ── Skip children that are inside an already-emitted accordion item ──
        if _find_accordion_parent(el) and id(_find_accordion_parent(el)) in seen_nodes:
            continue

        # ── Text-editor divs: may contain bare text nodes (no <p> wrapper) ──
        if _has_class(el, "elementor-widget-text-editor"):
            txt = el.get_text(" ", strip=True)
            if txt and len(txt) >= 5:
                key = txt.lower()
                if key not in seen:
                    seen.add(key)
                    text_items.append(("p", txt))
            continue

        # ── Button widgets: emit per-widget so duplicate labels on different buttons
        #    (e.g. multiple "LEARN MORE" cards in a carousel) are each captured.
        #    Use the widget's data-id as the dedup key instead of the button text.
        #    Also capture the href so it can be shown below the button in the docx. ──
        if _has_class(el, "elementor-widget-button"):
            btn_span = el.select_one(".elementor-button-text")
            if btn_span:
                btn_text = btn_span.get_text(" ", strip=True)
                widget_id = el.get("data-id", "")
                dedup_key = f"btn::{widget_id}" if widget_id else btn_text.lower()
                btn_link = el.select_one("a.elementor-button")
                href = btn_link.get("href", "").strip() if btn_link else ""
                if btn_text and dedup_key not in seen:
                    seen.add(dedup_key)
                    text_items.append(("button", btn_text, href))
                # Always mark button text and href as seen so the standard walker
                # below doesn't re-emit the inner <a>/<button> or href as plain text
                if btn_text:
                    seen.add(btn_text.lower())
                if href:
                    seen.add(href.lower())
            continue

        # ── Standard content tags ──
        if el.name not in ("h1","h2","h3","h4","h5","p","li","a","button","span"):
            continue

        # Skip wrapper spans (contain child tags) — only keep leaf spans
        if el.name == "span" and el.find(True):
            continue

        txt = el.get_text(" ", strip=True)
        if not txt or len(txt) < 5:
            continue

        key = txt.lower()
        if key in seen:
            continue
        seen.add(key)

        text_items.append((el.name, txt))

    # JSON-LD FAQ fallback — appends at end only if questions weren't already found in DOM
    for script in soup.find_all("script", {"type": "application/ld+json"}):
        try:
            data = json.loads(script.string or "")
            if isinstance(data, dict):
                data = [data]
            for obj in data:
                if obj.get("@type") == "FAQPage":
                    for entity in obj.get("mainEntity", []):
                        q = entity.get("name", "")
                        a_obj = entity.get("acceptedAnswer", {})
                        a = a_obj.get("text", "") if isinstance(a_obj, dict) else ""
                        if q and a and q.lower() not in seen:
                            seen.add(q.lower())
                            a_clean = BeautifulSoup(a, "lxml").get_text(" ", strip=True)
                            text_items.append(("h3", q))
                            text_items.append(("p", a_clean))
        except Exception:
            pass

    return text_items

# ==============================
# SCRAPE PAGE
# ==============================

def scrape_page(url):
    global driver
    for _ in range(MAX_RETRIES):
        try:
            driver.get(url)
            time.sleep(WAIT_TIME)

            smart_scroll()                 # initial lazy-load trigger
            expand_dynamic_content()       # accordion-only clicks, 2 rounds
            force_reveal_hidden_content()  # JS force-show all hidden panels
            activate_sliders()             # 20-click carousel cycling + force-reveal
            smart_scroll()                 # catch newly-visible lazy items

            soup = BeautifulSoup(driver.page_source, "lxml")
            title = soup.title.string.strip() if soup.title else url
            content_blocks = clean_text_blocks(soup)

            return title, content_blocks

        except Exception as e:
            print(f"Retry for {url}: {e}")
            # Tab crash: restart the browser and try again
            if "tab crashed" in str(e).lower() or "session" in str(e).lower():
                print("  Browser crash detected — restarting Chrome...")
                try:
                    driver.quit()
                except Exception:
                    pass
                time.sleep(2)
                driver = make_driver()
            else:
                time.sleep(2)

    return url, []

# ==============================
# DOCX BUILDER
# ==============================

def build_docx(company_name, pages_data, output_path):

    doc = Document()

    title_para = doc.add_paragraph()
    run = title_para.add_run(f"Content - {company_name}")
    run.bold = True
    run.font.size = Pt(20)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Page Title"
    table.rows[0].cells[1].text = "Source URL"

    for page in pages_data:
        row = table.add_row().cells
        row[0].text = page["title"]
        row[1].text = page["url"]

    doc.add_page_break()

    for page in pages_data:
        doc.add_heading(page["title"], level=1)
        doc.add_paragraph(f"URL: {page['url']}")
        doc.add_paragraph("")

        for item in page["content"]:
            tag, text = item[0], item[1]
            href = item[2] if len(item) > 2 else ""

            if tag.startswith("h"):
                level = min(int(tag[1]), 4)
                doc.add_heading(text, level=level)
                doc.add_paragraph("")
            elif tag == "button":
                # Bold button label
                para = doc.add_paragraph()
                run = para.add_run(f"[ {text} ]")
                run.bold = True
                # Link on the line below, in smaller italic text
                if href:
                    link_para = doc.add_paragraph()
                    link_run = link_para.add_run(href)
                    link_run.italic = True
                    link_run.font.size = Pt(9)
            else:
                para = doc.add_paragraph(text)
                if any(p in text.lower() for p in CTA_PHRASES):
                    para.runs[0].bold = True

        doc.add_page_break()

    doc.save(output_path)

# ==============================
# MAIN
# ==============================

def main():

    sitemap_urls = find_sitemap_urls(BASE_URL)
    if not sitemap_urls:
        print("No sitemap found.")
        return

    all_urls = set()
    for sm in sitemap_urls:
        parse_sitemap(sm, all_urls)

    normal_pages = []
    low_priority_pages = []

    for i, url in enumerate(sorted(all_urls), 1):
        print(f"[{i}/{len(all_urls)}] Scraping {url}")

        title, content = scrape_page(url)

        page_obj = {
            "url": url,
            "title": title,
            "content": content,
        }

        if any(p in url.lower() for p in LOW_PRIORITY_PATTERNS):
            low_priority_pages.append(page_obj)
        else:
            normal_pages.append(page_obj)

    pages_data = normal_pages + low_priority_pages

    company_name = urlparse(BASE_URL).netloc.replace("www.", "")
    safe_name = company_name.replace(".", "_")
    output_path = f"{safe_name}_site_content.docx"

    print("Building DOCX...")
    build_docx(company_name, pages_data, output_path)

    driver.quit()
    print(f"Done! Saved to {output_path}")


if __name__ == "__main__":
    main()