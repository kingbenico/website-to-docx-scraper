"""
BALANCED PRO EXTRACTOR v3 — importable module
Refactored from bulk_site_to_docx_selenium.py for use as a library.
Driver is created inside run_scrape() — no module-level side effects.
"""

import ctypes
import json
import os
import queue
import re
import sys
import tempfile
import threading
import time
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
import requests
from urllib.parse import urljoin, urlparse, unquote
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import TimeoutException, WebDriverException
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==============================
# CONSTANTS
# ==============================

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7)"
}

LOW_PRIORITY_PATTERNS = [
    "/category/",
    "/tag/",
    "/author/",
    "/hello-world",
]

CTA_PHRASES = [
    "call now",
    "get in touch",
    "schedule",
    "request",
    "learn more",
    "submit",
    "free inspection",
    "contact us",
]

# ==============================
# SELENIUM SETUP
# ==============================

def make_driver():
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--disable-extensions")
    chrome_options.add_argument("--disable-software-rasterizer")
    chrome_options.add_argument("--js-flags=--max-old-space-size=512")
    chrome_options.add_argument("--disable-background-networking")
    chrome_options.add_argument("--disable-default-apps")
    chrome_options.add_argument("--disable-sync")
    chrome_options.add_argument("--metrics-recording-only")
    chrome_options.add_argument("--mute-audio")
    chrome_options.add_argument("--no-first-run")
    chrome_options.add_argument("--safebrowsing-disable-auto-update")
    # Unique throwaway profile per run. A shared/persistent --user-data-dir can
    # be left locked between runs on Render, which makes Chrome start in a
    # degraded state where driver.get() silently fails — breaking the sitemap
    # browser fallback.
    chrome_options.add_argument(f"--user-data-dir={tempfile.mkdtemp(prefix='chrome-')}")
    driver = webdriver.Chrome(options=chrome_options)
    # A hung resource on one page shouldn't block the whole scrape forever.
    driver.set_page_load_timeout(30)
    driver.set_script_timeout(30)
    return driver


class DriverPool:
    """
    Fixed-size pool of Chrome drivers for parallel page scraping. Drivers are
    created lazily on first checkout, so a single-page job never pays to start
    N browsers. Uses a queue.Queue (not threading.local) because scrape_page
    can REPLACE a driver after a crash — a Queue makes returning that
    replacement to the pool straightforward and correct; thread-local storage
    would silently leak the old driver if the executor recycles threads.
    """

    def __init__(self, size, log=print):
        self.size = max(1, size)
        self.log = log
        self._q = queue.Queue()
        self._all = []
        self._lock = threading.Lock()
        self._created = 0

    def seed(self, drv):
        """Add an already-created driver to the pool instead of starting a new
        one — used to hand off the sitemap-discovery driver so we never
        briefly run N+1 Chromes."""
        with self._lock:
            self._all.append(drv)
            self._created += 1
        self._q.put(drv)

    def acquire(self):
        try:
            return self._q.get_nowait()
        except queue.Empty:
            pass
        with self._lock:
            if self._created < self.size:
                self._created += 1
                make_new = True
            else:
                make_new = False
        if make_new:
            self.log(f"  [pool] starting Chrome #{self._created}/{self.size}")
            drv = make_driver()
            with self._lock:
                self._all.append(drv)
            return drv
        return self._q.get()  # all drivers checked out — wait for one to free up

    def release(self, drv):
        self._q.put(drv)

    def replace(self, old, new):
        """scrape_page returned a different driver than it was given (crash restart)."""
        with self._lock:
            if old in self._all:
                self._all.remove(old)
            self._all.append(new)

    def shutdown(self):
        with self._lock:
            drivers, self._all = self._all, []
        for d in drivers:
            try:
                d.quit()
            except Exception:
                pass


def _total_ram_mb():
    """Best-effort total system RAM in MB, without adding a psutil dependency
    (kept out to avoid bloating the PyInstaller desktop build)."""
    try:
        if sys.platform == "win32":
            class MEMORYSTATUSEX(ctypes.Structure):
                _fields_ = [
                    ("dwLength", ctypes.c_ulong),
                    ("dwMemoryLoad", ctypes.c_ulong),
                    ("ullTotalPhys", ctypes.c_ulonglong),
                    ("ullAvailPhys", ctypes.c_ulonglong),
                    ("ullTotalPageFile", ctypes.c_ulonglong),
                    ("ullAvailPageFile", ctypes.c_ulonglong),
                    ("ullTotalVirtual", ctypes.c_ulonglong),
                    ("ullAvailVirtual", ctypes.c_ulonglong),
                    ("sullAvailExtendedVirtual", ctypes.c_ulonglong),
                ]
            stat = MEMORYSTATUSEX()
            stat.dwLength = ctypes.sizeof(MEMORYSTATUSEX)
            ctypes.windll.kernel32.GlobalMemoryStatusEx(ctypes.byref(stat))
            return stat.ullTotalPhys / (1024 * 1024)
        else:
            page_size = os.sysconf("SC_PAGE_SIZE")
            page_count = os.sysconf("SC_PHYS_PAGES")
            return (page_size * page_count) / (1024 * 1024)
    except Exception:
        return 2048  # safe fallback: assume a modest 2GB


def _default_workers():
    """
    How many Chrome instances to run in parallel. Priority order:
    1. SCRAPER_WORKERS env var (explicit override, e.g. set to 1 on Render).
    2. Derived from total RAM (~700MB budget per Chrome instance) and CPU
       count, capped at 4. Returns 1 below 1GB RAM so this never OOMs
       Render's 512MB free tier by defaulting to multiple browsers there.
    """
    env = os.environ.get("SCRAPER_WORKERS")
    if env:
        try:
            return max(1, min(8, int(env)))
        except ValueError:
            pass
    total_mb = _total_ram_mb()
    if total_mb < 1024:
        return 1
    cpu = os.cpu_count() or 2
    return max(1, min(4, cpu // 2, int(total_mb // 700)))

# ==============================
# SITEMAP
# ==============================

def find_sitemap_urls(base_url, driver=None, log=print):
    """
    Look for a sitemap at common paths. Tries a plain HTTP request first (fast,
    no browser needed). If that finds nothing — e.g. a bot-protection layer
    like Cloudflare is blocking the hosting provider's IP with a 403/challenge
    page, which plain `requests` can't get past — and a Selenium `driver` was
    passed in, retry each candidate through the real browser instead, since a
    genuine browser TLS/JS fingerprint is far less likely to be blocked.
    """
    candidates = [
        "/sitemap.xml",
        "/sitemap_index.xml",
        "/wp-sitemap.xml",
    ]

    found_urls = []

    for path in candidates:
        url = urljoin(base_url, path)
        try:
            r = requests.get(url, timeout=20, headers=HEADERS)
            if r.status_code == 200:
                txt = r.text.lower()
                if "<urlset" in txt or "<sitemapindex" in txt:
                    found_urls.append(url)
                else:
                    log(f"  [sitemap] {url} -> HTTP 200 but not a sitemap (likely a challenge/error page)")
            else:
                log(f"  [sitemap] {url} -> HTTP {r.status_code}")
        except Exception as e:
            log(f"  [sitemap] {url} -> request failed: {e}")

    if found_urls:
        return list(set(found_urls))

    if driver is not None:
        log("  [sitemap] Plain HTTP requests found nothing — retrying via browser...")
        # An in-page XHR is subject to same-origin policy and returns nothing
        # if the browser hasn't loaded a page on this origin yet. If loading the
        # homepage fails there's no point firing XHRs — they'd all return '' —
        # so surface the real error and bail out of the browser fallback.
        try:
            driver.get(base_url)
        except Exception as e:
            log(f"  [sitemap] Could not load {base_url} in browser: {e} — skipping browser fallback")
            return list(set(found_urls))
        for path in candidates:
            url = urljoin(base_url, path)
            try:
                txt = _browser_fetch_text(driver, url)
                if txt and ("<urlset" in txt.lower() or "<sitemapindex" in txt.lower()):
                    found_urls.append(url)
                    log(f"  [sitemap] {url} -> found via browser")
                else:
                    snippet = (txt or "")[:120].replace("\n", " ")
                    log(f"  [sitemap] {url} -> browser fetch returned no sitemap content "
                        f"(len={len(txt or '')}): {snippet!r}")
            except Exception as e:
                log(f"  [sitemap] {url} -> browser fetch failed: {e}")

    return list(set(found_urls))


def _browser_fetch_text(driver, url):
    """
    Fetch a URL's raw response body using an in-page XMLHttpRequest, rather
    than driver.get(). Many WordPress sitemaps declare an xml-stylesheet
    (e.g. Rank Math's main-sitemap.xsl) that browsers apply when navigating
    directly to the URL, turning driver.page_source into XSLT-rendered HTML
    with no literal <urlset>/<sitemapindex> tags. A same-origin XHR returns
    the actual raw XML instead, while still riding on the browser's TLS/JS
    fingerprint to get past bot-protection that blocks plain `requests` calls.
    """
    return driver.execute_async_script(
        """
        var url = arguments[0];
        var callback = arguments[arguments.length - 1];
        var xhr = new XMLHttpRequest();
        xhr.open('GET', url);
        xhr.onload = function() { callback(xhr.responseText); };
        xhr.onerror = function() { callback(''); };
        xhr.send();
        """,
        url,
    )


def parse_sitemap(url, collected, driver=None, log=print):
    try:
        r = requests.get(url, timeout=30, headers=HEADERS)
        txt = r.text
        if r.status_code != 200 or ("<urlset" not in txt.lower() and "<sitemapindex" not in txt.lower()):
            if driver is not None:
                log(f"  [sitemap] {url} -> HTTP {r.status_code} via requests, retrying via browser...")
                txt = _browser_fetch_text(driver, url)
            else:
                log(f"  [sitemap] {url} -> HTTP {r.status_code}, unusable and no browser fallback available")
                return
        soup = BeautifulSoup(txt, "xml")

        sitemap_tags = soup.find_all("sitemap")
        if sitemap_tags:
            for sm in sitemap_tags:
                parse_sitemap(sm.loc.text.strip(), collected, driver=driver, log=log)
            return

        for url_tag in soup.find_all("url"):
            collected.add(url_tag.loc.text.strip())

    except Exception as e:
        log(f"Sitemap parse error: {e}")

# ==============================
# PAGE INTERACTION
# ==============================

def smart_scroll(driver, max_steps=30, settle=0.15, stable_needed=2):
    """
    Scroll to the bottom, continuing while scrollHeight keeps growing (lazy
    loaders extend it as we go). Exits once height is stable AND we're at the
    bottom for `stable_needed` consecutive checks, instead of a fixed step
    count — most pages finish in a handful of steps, and unlike the old fixed
    14-step version, long pages are no longer truncated at ~15k px.
    """
    try:
        stable = 0
        last_height = driver.execute_script("return document.body.scrollHeight;")
        for _ in range(max_steps):
            driver.execute_script("window.scrollBy(0, window.innerHeight);")
            time.sleep(settle)
            height, y, inner_h = driver.execute_script(
                "return [document.body.scrollHeight, window.pageYOffset, window.innerHeight];"
            )
            at_bottom = (y + inner_h) >= (height - 5)
            if height == last_height and at_bottom:
                stable += 1
                if stable >= stable_needed:
                    break
            else:
                stable = 0
            last_height = height
    except Exception:
        pass


def wait_dom_stable(driver, quiet=0.4, timeout=3.0):
    """
    Poll the rendered DOM size until it stops changing for `quiet` seconds, or
    `timeout` elapses. Used as a short settle period after JS-driven widgets
    (sliders, accordions) finish mutating the page.
    """
    deadline = time.time() + timeout
    last_len = -1
    stable_since = None
    while time.time() < deadline:
        try:
            n = driver.execute_script("return document.body.innerHTML.length;")
        except Exception:
            return
        if n == last_len:
            if stable_since is None:
                stable_since = time.time()
            elif time.time() - stable_since >= quiet:
                return
        else:
            stable_since = None
            last_len = n
        time.sleep(0.1)


def wait_ready(driver, timeout=10, min_wait=0.3):
    """
    Replace a blind fixed sleep with an actual readiness check: wait for
    document.readyState to reach "complete", then (if jQuery is present) for
    jQuery.active to reach 0, then a small floor for post-load JS injection.
    """
    try:
        WebDriverWait(driver, timeout).until(
            lambda d: d.execute_script("return document.readyState;") == "complete"
        )
    except Exception:
        pass
    try:
        WebDriverWait(driver, min(timeout, 5)).until(
            lambda d: d.execute_script(
                "return (window.jQuery === undefined) || (jQuery.active === 0);"
            )
        )
    except Exception:
        pass
    time.sleep(min_wait)


def force_reveal_hidden_content(driver):
    try:
        driver.execute_script("""
            // Pass 1: Accordion / FAQ / Collapse panels
            var accordionSelectors = [
                '.elementor-tab-content',
                '.elementor-accordion-content',
                '.elementor-toggle-content',
                '[class*="accordion-content"]',
                '[class*="faq-answer"]',
                '[class*="faq-content"]',
                '[class*="collapse"]',
                '.e-n-accordion-item__body'
            ];
            accordionSelectors.forEach(function(sel) {
                document.querySelectorAll(sel).forEach(function(el) {
                    el.style.display = 'block';
                    el.style.height = 'auto';
                    el.style.maxHeight = 'none';
                    el.style.overflow = 'visible';
                    el.style.visibility = 'visible';
                    el.style.opacity = '1';
                    el.removeAttribute('aria-hidden');
                    el.removeAttribute('hidden');
                });
            });

            // Pass 2: Tab panels
            var tabSelectors = [
                '[role="tabpanel"]',
                '.tab-pane',
                '.ui-tabs-panel',
                '.elementor-tab-content'
            ];
            tabSelectors.forEach(function(sel) {
                document.querySelectorAll(sel).forEach(function(el) {
                    el.style.display = 'block';
                    el.classList.remove('hidden', 'inactive');
                });
            });

            // Pass 3: Force ALL swiper/carousel slides fully visible (including clones)
            document.querySelectorAll('.swiper-slide, .slick-slide, .owl-item, .splide__slide, .elementor-carousel-item').forEach(function(el) {
                el.style.display = 'block';
                el.style.visibility = 'visible';
                el.style.opacity = '1';
                el.removeAttribute('aria-hidden');
                el.removeAttribute('hidden');
                el.removeAttribute('inert');
            });

            // Pass 4: Generic inline-hidden content-bearing elements
            var excludeTags = ['NAV', 'HEADER', 'FOOTER'];
            var excludeClasses = ['modal', 'overlay', 'popup', 'cookie', 'banner'];
            var hiddenEls = document.querySelectorAll(
                '[style*="display: none"], [style*="display:none"], [hidden], [aria-hidden="true"], [inert]'
            );
            hiddenEls.forEach(function(el) {
                if (excludeTags.indexOf(el.tagName) !== -1) return;
                var cls = (el.className && typeof el.className === 'string') ? el.className : '';
                for (var i = 0; i < excludeClasses.length; i++) {
                    if (cls.indexOf(excludeClasses[i]) !== -1) return;
                }
                el.style.display = 'block';
                el.style.height = 'auto';
                el.style.maxHeight = 'none';
                el.style.overflow = 'visible';
                el.style.visibility = 'visible';
                el.style.opacity = '1';
                el.removeAttribute('aria-hidden');
                el.removeAttribute('hidden');
                el.removeAttribute('inert');
            });
        """)
        time.sleep(0.3)
    except Exception:
        pass


def expand_dynamic_content(driver):
    """
    Open accordions/tabs by clicking. A SINGLE in-page pass: one round-trip
    total instead of 2 execute_script calls per element, and — critically —
    only one pass, since a second click pass toggles already-open panels shut
    again (accordions/<summary> are ON/OFF toggles, not idempotent "open"
    actions).
    """
    try:
        driver.execute_script(
            """
            var sels = ['.elementor-tab-title', '.elementor-accordion-title',
                        '.elementor-toggle-title', "[role='tab']", 'summary'];
            sels.forEach(function(sel) {
                document.querySelectorAll(sel).forEach(function(el) {
                    try {
                        // <details><summary> toggles via the .open property,
                        // not a click (a click would flip it shut).
                        if (el.tagName === 'SUMMARY' && el.parentElement) {
                            el.parentElement.open = true;
                            return;
                        }
                        var expanded = el.getAttribute('aria-expanded');
                        var cls = (el.className && typeof el.className === 'string') ? el.className : '';
                        if (expanded === 'true' || /active|open|expanded/i.test(cls)) {
                            return;  // already open — clicking would close it
                        }
                        el.click();
                    } catch (e) {}
                });
            });
            """
        )
        time.sleep(0.4)  # one settle for all animations, was 0.25s * N elements
    except Exception:
        pass


SLIDE_SELECTOR = (
    ".swiper-slide, .slick-slide, .owl-item, .splide__slide, .elementor-carousel-item"
)


def activate_sliders(driver, base_clicks=4, extra_clicks=8, click_delay=0.12, max_arrows=6):
    """
    Nudge carousels a few times so lazily-created slides get instantiated.
    force_reveal_hidden_content() already CSS-forces every EXISTING slide to
    display:block/visible/opacity:1, so for plain text extraction full 20x
    cycling is unnecessary — all slides already sit in the DOM. The one case
    that still needs clicking is a carousel that creates slide elements on
    demand; we detect that by counting slides before/after a small number of
    clicks and only escalate if the count actually grew.
    """
    try:
        before = driver.execute_script(
            f"return document.querySelectorAll('{SLIDE_SELECTOR}').length;"
        )
        if not before:
            return

        arrows = driver.find_elements(
            By.CSS_SELECTOR,
            ".swiper-button-next, .slick-next, .elementor-swiper-button-next, "
            ".owl-next, .splide__arrow--next, [aria-label='Next']",
        )[:max_arrows]

        def click_arrows(times):
            for arrow in arrows:
                for _ in range(times):
                    try:
                        driver.execute_script("arguments[0].click();", arrow)
                        time.sleep(click_delay)
                    except Exception:
                        break

        click_arrows(base_clicks)

        after = driver.execute_script(
            f"return document.querySelectorAll('{SLIDE_SELECTOR}').length;"
        )
        if after > before:
            # Slides are being created lazily — keep going a bit longer.
            click_arrows(extra_clicks)

        force_reveal_hidden_content(driver)

    except Exception:
        pass

# ==============================
# FAQ SPECIAL HANDLER
# ==============================

def extract_elementor_faq(soup):
    faq_blocks = []
    seen_questions = set()

    for item in soup.select(".elementor-accordion-item"):
        q = item.select_one(".elementor-tab-title")
        a = item.select_one(".elementor-tab-content")
        if q and a:
            question = q.get_text(" ", strip=True)
            answer = a.get_text(" ", strip=True)
            if question and answer and question.lower() not in seen_questions:
                seen_questions.add(question.lower())
                faq_blocks.append(("h3", question))
                faq_blocks.append(("p", answer))

    for item in soup.select(".e-n-accordion-item"):
        q = item.select_one(".e-n-accordion-item-title-text") or item.select_one("summary")
        a = item.select_one(".e-n-accordion-item__body")
        if q and a:
            question = q.get_text(" ", strip=True)
            answer = a.get_text(" ", strip=True)
            if question and answer and question.lower() not in seen_questions:
                seen_questions.add(question.lower())
                faq_blocks.append(("h3", question))
                faq_blocks.append(("p", answer))

    for detail in soup.find_all("details"):
        summary_tag = detail.find("summary")
        if not summary_tag:
            continue
        question = summary_tag.get_text(" ", strip=True)
        answer_parts = []
        for child in detail.children:
            if child == summary_tag:
                continue
            if hasattr(child, "get_text"):
                part = child.get_text(" ", strip=True)
                if part:
                    answer_parts.append(part)
        answer = " ".join(answer_parts)
        if question and answer and question.lower() not in seen_questions:
            seen_questions.add(question.lower())
            faq_blocks.append(("h3", question))
            faq_blocks.append(("p", answer))

    for script in soup.find_all("script", {"type": "application/ld+json"}):
        try:
            data = json.loads(script.string or "")
            if isinstance(data, dict):
                data = [data]
            for obj in data:
                if obj.get("@type") == "FAQPage":
                    for entity in obj.get("mainEntity", []):
                        q = entity.get("name", "")
                        a_obj = entity.get("acceptedAnswer", {})
                        a = a_obj.get("text", "") if isinstance(a_obj, dict) else ""
                        if q and a and q.lower() not in seen_questions:
                            seen_questions.add(q.lower())
                            a_clean = BeautifulSoup(a, "lxml").get_text(" ", strip=True)
                            faq_blocks.append(("h3", q))
                            faq_blocks.append(("p", a_clean))
        except Exception:
            pass

    return faq_blocks

# ==============================
# BUSINESS INFO EXTRACTOR
# ==============================

PHONE_RE = re.compile(
    r"(\+?1[\s.-]?)?\(?\d{3}\)?[\s.-]\d{3}[\s.-]\d{4}\b"
)
EMAIL_RE = re.compile(
    r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+"
)
ADDRESS_HINT_RE = re.compile(
    r"\d{1,6}\s+[A-Za-z0-9.\s]{2,60}\b(Street|St|Avenue|Ave|Boulevard|Blvd|Road|Rd|Drive|Dr|Lane|Ln|Way|Court|Ct|Circle|Cir|Place|Pl|Suite|Ste|Highway|Hwy)\b[.,]?\s*[A-Za-z\s]{0,40},?\s*[A-Z]{2}\s*\d{5}(-\d{4})?",
    re.IGNORECASE,
)

EMAIL_EXCLUDE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".gif", ".svg", ".webp")


def _bump(info, key, value):
    """Record one occurrence of `value` in the info Counter at `key`."""
    if key == "phones" and value:
        # Collapse a repeated leading '+' (e.g. a buggy '++18564859091' from a
        # doubled tel: prefix) down to a single '+'. This does NOT merge distinct
        # formats — '(856) 485-9091' and '+18564859091' remain separate entries.
        value = re.sub(r"^\++", "+", value.strip())
    if value:
        info[key][value] += 1


def _extract_phones(soup, info):
    for a in soup.find_all("a", href=True):
        href = unquote(a["href"].strip())
        if href.lower().startswith("tel:"):
            phone = href[4:].strip()
            if phone:
                _bump(info, "phones", phone)

    body = soup.body if soup.body else soup
    text = body.get_text(" ", strip=True)
    for m in PHONE_RE.finditer(text):
        candidate = m.group(0).strip()
        digits = re.sub(r"\D", "", candidate)
        if len(digits) in (10, 11):
            _bump(info, "phones", candidate)


def extract_business_info(soup, url, info, skip_phones=False):
    """
    Scan a page's soup for business contact details (phone, email, address)
    and merge findings into the shared `info` dict. Also checks schema.org
    LocalBusiness/Organization JSON-LD, which WordPress sites commonly include.

    skip_phones: when True, phone extraction is skipped here — the caller is
    responsible for extracting phones separately from raw (non-JS) HTML, to
    avoid picking up JS-swapped call-tracking numbers.
    """

    # --- JSON-LD structured data (most reliable source) ---
    for script in soup.find_all("script", {"type": "application/ld+json"}):
        try:
            data = json.loads(script.string or "")
        except Exception:
            continue

        candidates = data if isinstance(data, list) else [data]
        for obj in candidates:
            if not isinstance(obj, dict):
                continue
            obj_type = obj.get("@type", "")
            if isinstance(obj_type, list):
                is_business = any(
                    t in ("LocalBusiness", "Organization", "Corporation")
                    or "Business" in str(t)
                    for t in obj_type
                )
            else:
                is_business = (
                    obj_type in ("LocalBusiness", "Organization", "Corporation")
                    or "Business" in str(obj_type)
                )
            if not is_business:
                continue

            if not info["name"] and obj.get("name"):
                info["name"] = obj["name"].strip()

            if not skip_phones:
                phone = obj.get("telephone")
                if phone:
                    _bump(info, "phones", phone.strip())

            email = obj.get("email")
            if email:
                _bump(info, "emails", email.strip())

            addr = obj.get("address")
            if isinstance(addr, dict):
                parts = [
                    addr.get("streetAddress", ""),
                    addr.get("addressLocality", ""),
                    addr.get("addressRegion", ""),
                    addr.get("postalCode", ""),
                ]
                full_addr = ", ".join(p.strip() for p in parts if p and p.strip())
                if full_addr:
                    _bump(info, "addresses", full_addr)
            elif isinstance(addr, str) and addr.strip():
                _bump(info, "addresses", addr.strip())

    # --- tel: link / plain-text phone scan (skipped when using raw-HTML mode) ---
    if not skip_phones:
        _extract_phones(soup, info)

    # --- mailto: links ---
    for a in soup.find_all("a", href=True):
        href = unquote(a["href"].strip())
        if href.lower().startswith("mailto:"):
            email = href[7:].split("?")[0].strip()
            if email:
                _bump(info, "emails", email)

    # --- Plain-text regex scan of visible body text (footer/header often hold this) ---
    body = soup.body if soup.body else soup
    text = body.get_text(" ", strip=True)

    for m in EMAIL_RE.finditer(text):
        candidate = m.group(0).strip().rstrip(".,;")
        if not candidate.lower().endswith(EMAIL_EXCLUDE_EXTENSIONS):
            _bump(info, "emails", candidate)

    for m in ADDRESS_HINT_RE.finditer(text):
        _bump(info, "addresses", m.group(0).strip())

    # --- Google Maps embed fallback ---
    # Embedded Maps widgets (<gmp-place-details-compact>, older iframe embeds) render the
    # address inside a closed shadow root, which is invisible to page_source / DOM scripting.
    # No address text can be recovered from that widget without calling the Places API, so
    # as a fallback we capture the "Get directions" / map link instead, giving the user a
    # clickable reference to the location even when the plain-text address isn't found.
    if not info["addresses"]:
        for a in soup.find_all("a", href=True):
            href = a["href"].strip()
            if "google.com/maps" in href.lower() or "maps.google.com" in href.lower():
                info["map_links"].add(href)


_thread_local = threading.local()


def _requests_session():
    """Thread-local requests.Session so connection pooling works safely once
    pages are fetched from multiple worker threads (Phase 2 parallelism)."""
    s = getattr(_thread_local, "session", None)
    if s is None:
        s = requests.Session()
        s.headers.update(HEADERS)
        _thread_local.session = s
    return s


def fetch_static_soup(url):
    """
    Fetch a page's raw HTML via plain HTTP (no browser, no JavaScript at all)
    and parse it. Returns a BeautifulSoup, or None on any failure. Callers
    that need the raw HTML for multiple purposes (static phone extraction AND
    the dynamic-phone body swap) should call this ONCE and share the result —
    fetching per-purpose was issuing up to 4 identical GETs per page.
    """
    try:
        r = _requests_session().get(url, timeout=20)
        if r.status_code != 200:
            return None
        return BeautifulSoup(r.text, "lxml")
    except Exception:
        return None


def extract_static_phones(static_soup, info):
    """
    Extract phone numbers from a page's raw (non-JS) HTML soup. Used when the
    user wants the original static phone number rather than one swapped in by
    a call-tracking script (e.g. CallRail, CallTrackingMetrics) that only runs
    after JS executes. `static_soup` is the result of fetch_static_soup().
    """
    if static_soup is None:
        return

    _extract_phones(static_soup, info)

    for script in static_soup.find_all("script", {"type": "application/ld+json"}):
        try:
            data = json.loads(script.string or "")
        except Exception:
            continue
        candidates = data if isinstance(data, list) else [data]
        for obj in candidates:
            if not isinstance(obj, dict):
                continue
            phone = obj.get("telephone")
            if phone:
                _bump(info, "phones", phone.strip())


def _collect_phone_strings(soup):
    """
    Return the set of raw phone-number strings found in `soup` (both tel: links
    and visible text). Used to compare the JS-rendered DOM against the raw HTML
    so we can tell which numbers were injected by a call-tracking script.
    """
    tmp = {"phones": Counter()}
    _extract_phones(soup, tmp)
    return set(tmp["phones"].keys())


def _compute_phone_swap(static_soup, rendered_soup, log):
    """
    Compare the raw-HTML (static) soup against the JS-rendered soup and decide
    what dynamic->static phone substitution, if any, should be applied to the
    page body. Pure — no HTTP, no block mutation.

    A number is "dynamic" if it appears in the JS-rendered DOM but its last-10
    digits are NOT present in the raw HTML. We only propose a swap when there
    is exactly ONE static number, so we never risk mangling a page that
    legitimately lists several different numbers.

    Returns (dynamic_strings, static_display) or (None, None) if no swap
    should happen.
    """
    def digits10(s):
        d = re.sub(r"\D", "", s)
        return d[-10:] if len(d) >= 10 else d

    if static_soup is None:
        return None, None

    static_strings = _collect_phone_strings(static_soup)
    static_keys = {digits10(s) for s in static_strings if digits10(s)}
    if len(static_keys) != 1:
        if len(static_keys) > 1:
            log("  [static-phones] multiple static numbers found — skipping body swap to avoid a wrong replacement")
        return None, None
    static_key = next(iter(static_keys))
    # Pick a canonical display string for that static number (the longest raw
    # form tends to be the most complete, e.g. with area code / punctuation).
    static_display = sorted(
        (s for s in static_strings if digits10(s) == static_key), key=len, reverse=True
    )[0]

    rendered_strings = _collect_phone_strings(rendered_soup)
    dynamic_strings = {s for s in rendered_strings if digits10(s) and digits10(s) != static_key}
    if not dynamic_strings:
        return None, None
    return dynamic_strings, static_display


def _apply_phone_swap(blocks, dynamic_strings, static_display):
    """
    Apply a precomputed dynamic->static swap to one block list, in place.

    `blocks` is a list of (tag, text) tuples, or (tag, text, href) for
    buttons — both the visible text (index 1) and a button's href (index 2,
    e.g. 'tel:+18564859091') can carry the dynamic number. Returns the number
    of blocks changed.
    """
    def swap(text):
        new_text = text
        for dyn in dynamic_strings:
            if dyn in new_text:
                new_text = new_text.replace(dyn, static_display)
        return new_text

    subs = 0
    new_blocks = []
    for block in blocks:
        block = list(block)
        changed = False
        for i in range(1, len(block)):
            if isinstance(block[i], str):
                swapped = swap(block[i])
                if swapped != block[i]:
                    block[i] = swapped
                    changed = True
        if changed:
            subs += 1
        new_blocks.append(tuple(block))
    if subs:
        blocks[:] = new_blocks
    return subs


# ==============================
# CLEAN TEXT
# ==============================

def extract_text_blocks(root, strip_forms=True):
    """
    Walk a subtree (page body, or a standalone header/footer fragment) and
    extract the same heading/paragraph/button/accordion structure that
    clean_text_blocks() produces for the main page body. Shared so header
    and footer extraction stay visually consistent with page content.
    """
    if strip_forms:
        for tag in root(["input", "textarea", "select", "option", "label"]):
            tag.decompose()

    text_items = []
    seen = set()
    seen_nodes = set()

    def _has_class(el, cls):
        return cls in (el.get("class") or [])

    def _find_accordion_parent(el):
        for parent in el.parents:
            if _has_class(parent, "elementor-accordion-item") or _has_class(parent, "e-n-accordion-item"):
                return parent
        return None

    for el in root.find_all(True, recursive=True):

        if _has_class(el, "elementor-accordion-item") or _has_class(el, "e-n-accordion-item"):
            acc_id = id(el)
            if acc_id not in seen_nodes:
                seen_nodes.add(acc_id)

                q_tag = el.select_one(".elementor-tab-title")
                a_tag = el.select_one(".elementor-tab-content")

                if not q_tag:
                    q_tag = el.select_one(".e-n-accordion-item-title-text")
                if not a_tag:
                    a_tag = el.find("div", attrs={"role": "region"})

                if q_tag and a_tag:
                    q_text = q_tag.get_text(" ", strip=True)
                    a_text = a_tag.get_text(" ", strip=True)
                    if q_text and q_text.lower() not in seen:
                        seen.add(q_text.lower())
                        text_items.append(("h3", q_text))
                    if a_text and a_text.lower() not in seen:
                        seen.add(a_text.lower())
                        text_items.append(("p", a_text))
            continue

        if _find_accordion_parent(el) and id(_find_accordion_parent(el)) in seen_nodes:
            continue

        if _has_class(el, "elementor-widget-text-editor"):
            txt = el.get_text(" ", strip=True)
            if txt and len(txt) >= 5:
                key = txt.lower()
                if key not in seen:
                    seen.add(key)
                    text_items.append(("p", txt))
            continue

        if _has_class(el, "elementor-widget-button"):
            btn_span = el.select_one(".elementor-button-text")
            if btn_span:
                btn_text = btn_span.get_text(" ", strip=True)
                widget_id = el.get("data-id", "")
                dedup_key = f"btn::{widget_id}" if widget_id else btn_text.lower()
                btn_link = el.select_one("a.elementor-button")
                href = btn_link.get("href", "").strip() if btn_link else ""
                if btn_text and dedup_key not in seen:
                    seen.add(dedup_key)
                    text_items.append(("button", btn_text, href))
                if btn_text:
                    seen.add(btn_text.lower())
                if href:
                    seen.add(href.lower())
            continue

        if el.name not in ("h1","h2","h3","h4","h5","p","li","a","button","span"):
            continue

        if el.name == "span" and el.find(True):
            continue

        txt = el.get_text(" ", strip=True)
        if not txt or len(txt) < 5:
            continue

        key = txt.lower()
        if key in seen:
            continue
        seen.add(key)

        text_items.append((el.name, txt))

    return text_items


def clean_text_blocks(soup):

    for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "svg"]):
        tag.decompose()

    root = soup.body if soup.body else soup
    text_items = extract_text_blocks(root)

    for script in soup.find_all("script", {"type": "application/ld+json"}):
        try:
            data = json.loads(script.string or "")
            if isinstance(data, dict):
                data = [data]
            for obj in data:
                if obj.get("@type") == "FAQPage":
                    for entity in obj.get("mainEntity", []):
                        q = entity.get("name", "")
                        a_obj = entity.get("acceptedAnswer", {})
                        a = a_obj.get("text", "") if isinstance(a_obj, dict) else ""
                        if q and a and q.lower() not in {i[1].lower() for i in text_items}:
                            a_clean = BeautifulSoup(a, "lxml").get_text(" ", strip=True)
                            text_items.append(("h3", q))
                            text_items.append(("p", a_clean))
        except Exception:
            pass

    return text_items

# ==============================
# SCRAPE PAGE
# ==============================

def scrape_page(driver, url, wait_time, max_retries, business_info=None, static_phones=False,
                 include_header=False, include_footer=False, log=print):
    """
    Navigate to `url`, let JS-driven content render, and extract its text.

    The retry loop only covers navigation/rendering (driver.get + widget
    interaction) — those can be transient (a slow load, a crashed tab) and are
    worth retrying. Parsing/extraction runs exactly once afterward: it's
    deterministic, so retrying it just re-runs the same interaction cost for
    the same result.
    """
    page_source = None
    for attempt in range(max_retries):
        try:
            try:
                driver.get(url)
            except TimeoutException:
                log(f"  [timeout] {url} exceeded page load timeout — extracting what rendered")
                try:
                    driver.execute_script("window.stop();")
                except Exception:
                    pass

            wait_ready(driver, timeout=max(10, wait_time * 3), min_wait=min(wait_time, 0.3))
            smart_scroll(driver)
            expand_dynamic_content(driver)
            force_reveal_hidden_content(driver)
            activate_sliders(driver)
            wait_dom_stable(driver)

            page_source = driver.page_source
            break

        except WebDriverException as e:
            msg = str(e).lower()
            log(f"  Retry {attempt + 1}/{max_retries} for {url}: {e}")
            if "tab crashed" in msg or "session" in msg or "disconnected" in msg:
                log("  Browser crash detected — restarting Chrome...")
                try:
                    driver.quit()
                except Exception:
                    pass
                time.sleep(2)
                driver = make_driver()
            else:
                time.sleep(1)

    if page_source is None:
        return url, [], driver, [], []

    # --- Parsing/extraction: deterministic, never retried ---
    try:
        soup = BeautifulSoup(page_source, "lxml")
        title = soup.title.string.strip() if soup.title else url

        # Fetch raw (non-JS) HTML once, up front, and share it for both static
        # phone extraction and the dynamic->static body swap below — avoids
        # issuing multiple identical GETs for the same URL.
        static_soup = fetch_static_soup(url) if static_phones else None

        # Business info (phone/email/address) often lives in header/footer/nav,
        # which clean_text_blocks() strips — extract it first, before that mutation.
        if business_info is not None:
            extract_business_info(soup, url, business_info, skip_phones=static_phones)
            if static_phones:
                extract_static_phones(static_soup, business_info)

        # Header/footer content, if requested — must run before clean_text_blocks()
        # decomposes <header>/<footer> tags out of the soup.
        header_blocks = []
        footer_blocks = []
        if include_header:
            for tag in soup.find_all("header"):
                header_blocks.extend(extract_text_blocks(tag, strip_forms=False))
        if include_footer:
            for tag in soup.find_all("footer"):
                footer_blocks.extend(extract_text_blocks(tag, strip_forms=False))

        content_blocks = clean_text_blocks(soup)

        # With static-phones mode on, the body above was extracted from the
        # JS-rendered DOM, so it still shows any call-tracking number that a
        # script swapped in. Rewrite those to the site's static number so the
        # document body matches the Business Summary.
        if static_phones and static_soup is not None:
            dynamic_strings, static_display = _compute_phone_swap(static_soup, soup, log=log)
            if dynamic_strings:
                total = sum(
                    _apply_phone_swap(blk, dynamic_strings, static_display)
                    for blk in (content_blocks, header_blocks, footer_blocks)
                )
                if total:
                    log(f"  [static-phones] replaced dynamic number(s) with static "
                        f"'{static_display}' in {total} block(s)")

        return title, content_blocks, driver, header_blocks, footer_blocks

    except Exception as e:
        log(f"  Parse error for {url} (not retrying): {e}")
        return url, [], driver, [], []

# ==============================
# DOCX BUILDER
# ==============================

def build_docx(company_name, pages_data, output_path, business_info=None):

    doc = Document()

    title_para = doc.add_paragraph()
    run = title_para.add_run(f"Content - {company_name}")
    run.bold = True
    run.font.size = Pt(20)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # --- Business Summary (before the table of contents) ---
    if business_info is not None:
        doc.add_heading("Business Summary", level=1)

        if business_info["name"]:
            p = doc.add_paragraph()
            p.add_run("Business Name: ").bold = True
            p.add_run(business_info["name"])

        # Build the contact-details table rows: (Type, Data, Occurrence).
        # phones/emails/addresses are Counters {value: times_seen}; sorted by
        # descending occurrence so the most-common value is first. Map links
        # carry no meaningful count, so their occurrence cell is left as "-".
        rows = []
        # Phones: keep every distinct format as its own row (no merging), so
        # variants like '+18564859091' and '(856) 485-9091' are all preserved.
        for label, key in (("Phone", "phones"), ("Email", "emails"), ("Address", "addresses")):
            counter = business_info.get(key) or {}
            for value, count in sorted(counter.items(), key=lambda kv: (-kv[1], kv[0])):
                rows.append((label, value, str(count)))
        for link in sorted(business_info.get("map_links", [])):
            rows.append(("Map Link", link, "-"))

        if rows:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for cell, heading in zip(hdr, ("Type", "Data", "Occurrence")):
                cell.text = ""
                run = cell.paragraphs[0].add_run(heading)
                run.bold = True
            for r_type, r_data, r_occ in rows:
                cells = table.add_row().cells
                cells[0].text = r_type
                cells[1].text = r_data
                cells[2].text = r_occ
        else:
            doc.add_paragraph("No business contact details were found on the site.")

        doc.add_page_break()

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Page Title"
    table.rows[0].cells[1].text = "Source URL"

    for page in pages_data:
        row = table.add_row().cells
        row[0].text = page["title"]
        row[1].text = page["url"]

    doc.add_page_break()

    for page in pages_data:
        doc.add_heading(page["title"], level=1)
        doc.add_paragraph(f"URL: {page['url']}")
        doc.add_paragraph("")

        if page.get("header"):
            doc.add_heading("Header", level=2)
            _add_content_blocks(doc, page["header"])
            doc.add_paragraph("")

        if page.get("footer"):
            doc.add_heading("Footer", level=2)
            _add_content_blocks(doc, page["footer"])
            doc.add_paragraph("")

        _add_content_blocks(doc, page["content"])

        doc.add_page_break()

    doc.save(output_path)


def _add_content_blocks(doc, blocks):
    for item in blocks:
        tag, text = item[0], item[1]
        href = item[2] if len(item) > 2 else ""

        if tag.startswith("h"):
            level = min(int(tag[1]), 4)
            doc.add_heading(text, level=level)
            doc.add_paragraph("")
        elif tag == "button":
            para = doc.add_paragraph()
            run = para.add_run(f"[ {text} ]")
            run.bold = True
            if href:
                link_para = doc.add_paragraph()
                link_run = link_para.add_run(href)
                link_run.italic = True
                link_run.font.size = Pt(9)
        else:
            para = doc.add_paragraph(text)
            if any(p in text.lower() for p in CTA_PHRASES):
                para.runs[0].bold = True

# ==============================
# PUBLIC ENTRYPOINT
# ==============================

def _empty_business_info():
    # phones/emails/addresses are Counters so we can report how many times
    # each value was seen across the scraped pages (occurrence column).
    # map_links stays a set — no occurrence count needed for it.
    return {
        "name": "",
        "phones": Counter(),
        "emails": Counter(),
        "addresses": Counter(),
        "map_links": set(),
    }


def _merge_business_info(target, parts, url_order):
    """
    Merge per-worker business_info dicts into `target`, walking `url_order` so
    'name' stays first-write-wins in the same order as the old sequential
    code — the merged result is deterministic and independent of which worker
    thread happened to finish first.
    """
    for url in url_order:
        info = parts.get(url)
        if not info:
            continue
        if not target["name"] and info["name"]:
            target["name"] = info["name"]
        target["phones"] += info["phones"]
        target["emails"] += info["emails"]
        target["addresses"] += info["addresses"]
        target["map_links"] |= info["map_links"]


def _scrape_one(pool, url, wait_time, max_retries, static_phones,
                 include_header, include_footer, log):
    """Runs on a pool worker thread: check out a driver, scrape one page with
    its own business_info, return the driver (or its crash-restart
    replacement) to the pool."""
    local_info = _empty_business_info()
    drv = pool.acquire()
    try:
        title, content, new_drv, header_blocks, footer_blocks = scrape_page(
            drv, url, wait_time, max_retries,
            business_info=local_info, static_phones=static_phones,
            include_header=include_header, include_footer=include_footer,
            log=log,
        )
        if new_drv is not drv:
            pool.replace(drv, new_drv)
            drv = new_drv
        page_obj = {
            "url": url,
            "title": title,
            "content": content,
            "header": header_blocks,
            "footer": footer_blocks,
        }
        return page_obj, local_info
    finally:
        pool.release(drv)


def run_scrape(base_url, wait_time=3, max_retries=3, progress_callback=None, static_phones=False,
               single_url=None, include_header=False, include_footer=False, workers=None):
    """
    Run the full scrape pipeline for base_url.
    Returns the absolute path to the generated .docx file.
    progress_callback(msg: str) is called with status updates if provided.

    static_phones: when True, phone numbers are read from each page's raw
    (non-JS) HTML instead of the Selenium-rendered DOM, so call-tracking
    scripts that swap in a dynamic number after page load are bypassed.

    single_url: when provided, skips sitemap discovery/crawling entirely and
    scrapes only the given URL(s). Accepts either a single URL string or a list
    of URLs. Useful for targeting specific pages instead of the whole site
    (e.g. a 447-sitemap site where you only want a handful of pages), or for
    quickly testing extraction changes (e.g. static_phones) against one page.

    include_header / include_footer: when True, each page's <header>/<footer>
    content is extracted separately (normally stripped as boilerplate) and
    rendered under a "Header"/"Footer" heading in the docx.

    workers: number of Chrome instances to run in parallel. Defaults to an
    auto-detected value based on available RAM (see _default_workers) —
    forced to 1 on low-memory hosts like Render's free tier.
    """
    _log_lock = threading.Lock()

    def log(msg):
        with _log_lock:
            print(msg)
            if progress_callback:
                progress_callback(msg)

    log("Starting browser...")
    driver = make_driver()

    try:
        if single_url:
            # single_url may be a single string or a list of URLs.
            target_urls = [single_url] if isinstance(single_url, str) else list(single_url)
            all_urls = set(target_urls)
            if len(all_urls) == 1:
                log(f"Single-page mode: scraping only {next(iter(all_urls))}")
            else:
                log(f"Specific-pages mode: scraping {len(all_urls)} page(s):")
                for u in target_urls:
                    log(f"  - {u}")
        else:
            log("Discovering sitemaps...")
            sitemap_urls = find_sitemap_urls(base_url, driver=driver, log=log)
            if not sitemap_urls:
                driver.quit()
                raise ValueError("No sitemap found at the provided URL.")

            all_urls = set()
            for sm in sitemap_urls:
                parse_sitemap(sm, all_urls, driver=driver, log=log)

        log(f"Found {len(all_urls)} URLs to scrape.")

        all_urls_list = sorted(all_urls)
        business_info = _empty_business_info()

        num_workers = workers or _default_workers()
        if len(all_urls_list) == 1:
            num_workers = 1  # never spin up a pool for a single page

        # The sitemap-discovery driver becomes the pool's first member instead
        # of being quit and re-created, so we never briefly run N+1 Chromes.
        pool = DriverPool(num_workers, log=log)
        pool.seed(driver)

        log(f"Scraping with {num_workers} parallel browser(s)...")

        results = {}
        infos = {}
        completed = 0
        progress_lock = threading.Lock()

        try:
            with ThreadPoolExecutor(max_workers=num_workers, thread_name_prefix="scrape") as ex:
                futures = {
                    ex.submit(
                        _scrape_one, pool, url, wait_time, max_retries,
                        static_phones, include_header, include_footer, log,
                    ): url
                    for url in all_urls_list
                }
                for fut in as_completed(futures):
                    url = futures[fut]
                    try:
                        page_obj, local_info = fut.result()
                    except Exception as e:
                        log(f"  [error] {url}: {e}")
                        page_obj = {"url": url, "title": url, "content": [], "header": [], "footer": []}
                        local_info = None
                    results[url] = page_obj
                    if local_info:
                        infos[url] = local_info
                    with progress_lock:
                        completed += 1
                        log(f"[{completed}/{len(all_urls_list)}] Done {url}")
        finally:
            pool.shutdown()
            driver = None  # already quit via pool.shutdown()

        _merge_business_info(business_info, infos, all_urls_list)

        # Rebuild deterministic order + normal/low-priority split from
        # all_urls_list, since results arrive in completion order, not URL order.
        normal_pages = []
        low_priority_pages = []
        for url in all_urls_list:
            page_obj = results.get(url)
            if page_obj is None:
                continue
            if any(p in url.lower() for p in LOW_PRIORITY_PATTERNS):
                low_priority_pages.append(page_obj)
            else:
                normal_pages.append(page_obj)

        pages_data = normal_pages + low_priority_pages

        company_name = urlparse(base_url).netloc.replace("www.", "")
        safe_name = company_name.replace(".", "_")
        # Use the OS temp dir (works on Windows too) instead of a hardcoded /tmp.
        output_path = os.path.join(tempfile.gettempdir(), f"{safe_name}_site_content.docx")

        log("Building DOCX...")
        build_docx(company_name, pages_data, output_path, business_info)

        log(f"Done! Saved to {output_path}")
        return output_path

    finally:
        if driver is not None:
            try:
                driver.quit()
            except Exception:
                pass
